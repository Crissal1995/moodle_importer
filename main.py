from collections import defaultdict
import pathlib
import docx
import re
import argparse

from xml_builder import generate_xmls_per_module
import model

uf_re = re.compile(r'(uf)[^a-z]*([^(]+)[(]?(\d+)?', re.I)
module_re = re.compile(r'(modulo)[^a-z]*([^(]+)[(]?(\d+)?', re.I)
question_re = re.compile(r"(.+)(?=\s+\(?slide)\s+\(?slide[s]?\s+((\d+[\s-]?)+)", re.I)
answer_re = re.compile(r'(a?\.?\s*(.+?)(?=b\.))(b\.\s+(.+?)(?=c\.))(c\.(\s+.+))', re.I)


def populate_document(doc_pathlib):
    counts = defaultdict(int)
    uf_before, module_before, question_before = [None] * 3
    model_doc = model.Document(doc_pathlib.stem)
    try:
        parsed_doc = docx.Document(str(doc_pathlib))
    except Exception as e:
        raise ValueError(str(e) + '.\nDocumento Word non valido oppure aperto e non salvato!')

    for paragraph in parsed_doc.paragraphs:
        text = paragraph.text.strip()
        # eliminiamo caratteri unicode che danno problemi
        text = text.replace('–', '-').replace('’', "'").replace('‘', "'").replace('“', '"').replace('”', '"')
        for char, repl in zip('aeiouAEIOU', 'àèìòùÀÈÌÒÙ'):  # fix lettere accentate
            text = text.replace("{}'".format(char), repl)
        text = re.sub(r'\s+', ' ', text).strip()  # sostituiamo più spazi con uno solo + strip
        test = text.lower()

        if not test:
            continue

        if test.startswith('uf'):
            groups = uf_re.search(text).groups()
            name, duration = groups[1:3]
            uf = model.Unity(counts['uf'], name, duration)
            model_doc.add_unity(uf)

            uf_before = uf
            module_before, question_before = [None] * 2

            counts['uf'] += 1
            for kw in ('module', 'question'):
                counts[kw] = 0

        elif test.startswith('modulo'):
            assert uf_before, 'Found module without unity!'

            groups = module_re.search(text).groups()
            name, duration = groups[1:3]
            module = model.Module(counts['module'], name, duration, unity=uf_before)
            uf_before.add_module(module)

            module_before = module
            question_before = None

            counts['module'] += 1
            counts['question'] = 0

        else:  # domanda, risposta, o entrambe!

            # se questo if è vero, ho sia domanda che risposte sulla stessa riga
            if question_re.match(text) and answer_re.match(question_re.sub("", text).strip()):
                question_before = parse_question(text, counts, module_before)
                parse_answer(text, question_before)

                # reset question
                question_before = None

            # altrimenti ho solo la domanda
            elif question_re.match(text):
                question_before = parse_question(text, counts, module_before)

            # altrimenti ho solo le risposte
            elif answer_re.match(text):
                parse_answer(text, question_before)
                question_before = None

            # altrimenti ho qualche errore
            else:
                raise ValueError(f"cannot match anything with: {text}")

        """
        elif test.startswith('risposta'):
            assert question_before, 'Found answer without question!'

            groups = answer_re.search(text).groups()
            is_correct, name = groups[1:3]
            is_correct = True if is_correct else False  # fix None

            answer = model.Answer(name, is_correct)
            question_before.add_answer(answer)

        elif test.replace('*', '').strip().startswith('slide'):
            assert question_before, 'Found jump to slide without question!'
            slides = set([int(el) for el in re.findall(r'(\d+)', test)])
            question_before.set_jump2slides(slides)
        """

    # sort questions based on jump2slide
    for uf in model_doc.unities:
        for module in uf.modules:
            module.sort_questions()

    # generate clusters
    for uf in model_doc.unities:
        for module in uf.modules:
            module.write_cluster()

    return model_doc


def parse_question(text, counts, module_before):
    assert module_before, 'Found question without module!'

    # parse question text and slides number
    qname, slides = question_re.match(text).group(1, 2)

    # create model Question
    question = model.Question(counts["question"], counts["global_question"], qname, module_before)

    # add Question to module before this question
    module_before.add_question(question)

    # parse slides from regex match
    slides = set([int(elem) for elem in slides.split("-")])
    question.set_jump2slides(slides)

    # increment counters
    counts['question'] += 1
    counts['global_question'] += 1

    return question


def parse_answer(text, question_before):
    assert question_before, 'Found answers without question!'

    # parse answer
    a, b, c = answer_re.match(question_re.sub("", text).strip()).group(2, 4, 6)

    for elem in (a, b, c):
        elem = elem.strip()  # remove start characters
        if elem.startswith("ok"):
            elem = elem[2:].strip()
            is_correct = True
        else:
            is_correct = False
        answer = model.Answer(elem, is_correct=is_correct)
        question_before.add_answer(answer)

    if not any(answer.is_correct for answer in question_before.answers):
        raise ValueError(f"No correct answer found for {question_before.name}")


if __name__ == '__main__':
    parser = argparse.ArgumentParser()
    parser.add_argument('folder', type=str, help='The folder of Word documents to parse')
    parser.add_argument('-no', '--no-ordered', default=False, action='store_true', help='Don\'t display questions ordered by (min) slide to jump')
    parser.add_argument('-ns', '--no-separated', default=False, action='store_true', help='Don\'t separate questions in groups')
    parser.add_argument('-nd', '--no-dump', default=False, action='store_true', help='Dump each question parsed to stdout instead of a textfile with same name of document')
    args = parser.parse_args()

    q_dir = pathlib.Path(args.folder)
    files = list(q_dir.glob('*.docx')) + list(q_dir.glob('*.doc'))

    for file in files:
        print('Start to work on', file)
        doc = populate_document(file)
        doc.check()
        filepath = None if args.no_dump else file.stem
        generate_xmls_per_module(doc)  # create safely folder 'generated'
        doc.print_questions(filepath=filepath, ordered=not args.no_ordered, separated=not args.no_separated)
        print('----------\n')
